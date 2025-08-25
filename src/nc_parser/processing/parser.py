from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from PIL import Image, ImageOps, ImageEnhance, ImageFilter
import numpy as np
import cv2
from pdfminer.high_level import extract_text as pdf_extract_text
from pytesseract import image_to_string as ocr_image_to_string
from bs4 import BeautifulSoup
import pdfplumber
from pdf2image import convert_from_path
import csv
import subprocess
import zipfile
import re
import json
import time
from charset_normalizer import from_path as detect_encoding_from_path
import ftfy
from pypdf import PdfReader
from nc_parser.core.settings import get_settings
from structlog import get_logger
from nc_parser.processing.captioning import caption_images_with_cache

try:
    from striprtf.striprtf import rtf_to_text  # type: ignore
except Exception:  # pragma: no cover
    rtf_to_text = None  # type: ignore

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    full_text: str
    pages: list[dict[str, Any]]
    timings_ms: dict[str, float] | None = None
    metrics: dict[str, Any] | None = None


def _read_text_file(path: Path) -> str:
    try:
        # Try UTF-8 first
        return path.read_text(encoding="utf-8", errors="strict")
    except Exception:
        try:
            best = detect_encoding_from_path(str(path)).best()
            if best is not None:
                return best.output_text(stripped=True)
        except Exception:
            pass
        # Fallback with ignore to always return something
        txt = path.read_text(encoding="utf-8", errors="ignore")
        return ftfy.fix_text(txt)


def _read_pdf_text(path: Path) -> str:
    return pdf_extract_text(str(path)) or ""


def _pdf_quick_sanity(path: Path) -> bool:
    """Cheap PDF sanity: open metadata and count pages; ensure EOF marker present.

    Returns True if file looks sane; False to fast-fail.
    """
    try:
        # EOF marker near the end
        with path.open("rb") as f:
            f.seek(max(0, path.stat().st_size - 2048))
            tail = f.read()
            if b"%%EOF" not in tail:
                return False
        # Pages accessible
        reader = PdfReader(str(path))
        if len(reader.pages) < 1:
            return False
        return True
    except Exception:
        return False


def _read_doc_binary_text(path: Path) -> str:
    """Read legacy .doc via antiword; fallback to empty on failure."""
    try:
        res = subprocess.run(
            ["antiword", "-w", "0", str(path)],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            timeout=30,
        )
        out = res.stdout.decode("utf-8", errors="ignore")
        return ftfy.fix_text(out).strip()
    except Exception:
        return ""


def _read_rtf_text(path: Path) -> str:
    try:
        if rtf_to_text is None:
            return ""
        raw = path.read_text(encoding="utf-8", errors="ignore")
        txt = rtf_to_text(raw)
        return ftfy.fix_text(txt).strip()
    except Exception:
        return ""


def _rtf_to_html(path: Path) -> str:
    """Convert RTF to HTML using unrtf; return empty string on failure.

    We use this to reconstruct tables with colspan/rowspan for better fidelity.
    """
    try:
        res = subprocess.run(
            ["unrtf", "--html", "--nopict", str(path)],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            timeout=30,
        )
        out = res.stdout.decode("utf-8", errors="ignore")
        # unrtf wraps HTML in extra info; try to extract BODY
        soup = BeautifulSoup(out, "lxml")
        body = soup.find("body")
        return str(body) if body else out
    except Exception:
        return ""


def _read_odt_text(path: Path) -> str:
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            with zf.open("content.xml") as f:
                xml = f.read().decode("utf-8", errors="ignore")
        # Extract visible text; lxml already available
        soup = BeautifulSoup(xml, "lxml")
        txt = soup.get_text("\n")
        return ftfy.fix_text(txt).strip()
    except Exception:
        return ""


def _normalize_output_text(text: str, *, drop_noise: bool = True) -> str:
    if not text:
        return ""
    t = ftfy.fix_text(text)
    t = t.replace("\r", "").replace("\xa0", " ")
    # Normalize unicode minus to hyphen
    t = t.replace("−", "-")
    # Collapse excessive spaces but preserve newlines
    t = re.sub(r"[ \t]+", " ", t)
    # Trim trailing spaces on lines
    t = "\n".join(line.rstrip() for line in t.splitlines())
    if drop_noise:
        # Strip UI/noise lines: short, mostly symbols
        def _is_noise(line: str) -> bool:
            if not line:
                return False
            # Remove common bullet-like chars for ratio check
            cleaned = re.sub(r"[A-Za-z0-9]", "", line)
            symbol_ratio = (len(cleaned) / max(1, len(line)))
            if len(line) <= 2 and symbol_ratio > 0.7:
                return True
            # Lines with almost no letters and lots of punctuation/symbols
            letters = len(re.findall(r"[A-Za-zА-Яа-я]", line))
            if letters == 0 and symbol_ratio > 0.7:
                return True
            # UI crumbs like isolated icons repeated
            if re.fullmatch(r"[•·©®™@©\-_=+~^`\|<>\(\)\[\]{}\\]+", line.strip()):
                return True
            return False
        lines = [ln for ln in t.splitlines() if not _is_noise(ln.strip())]
        t = "\n".join(lines)
    return t.strip()


def _extract_text_from_html(raw_html: str) -> str:
    try:
        soup = BeautifulSoup(raw_html, "lxml")
        for tag in soup(["script", "style", "noscript", "head", "title", "meta", "link", "svg"]):
            tag.decompose()
        for tag in soup(["nav", "header", "footer", "aside"]):
            tag.decompose()
        text = soup.get_text("\n")
        text = _normalize_output_text(text)
        # Deduplicate short lines repeated many times (menus, footers)
        lines = []
        seen: dict[str, int] = {}
        for ln in text.splitlines():
            key = ln.strip()
            if not key:
                lines.append(ln)
                continue
            count = seen.get(key, 0)
            if len(key) <= 64 and count >= 1:
                continue
            seen[key] = count + 1
            lines.append(ln)
        return "\n".join(lines).strip()
    except Exception:
        return _normalize_output_text(BeautifulSoup(raw_html, "lxml").get_text("\n"))


def _extract_html_tables_rows_from_html(raw_html: str) -> list[list[list[str]]]:
    # Legacy: keep simple extractor for backward compatibility
    rows_all: list[list[list[str]]] = []
    try:
        soup = BeautifulSoup(raw_html, "lxml")
        for tbl in soup.find_all("table"):
            rows: list[list[str]] = []
            for tr in tbl.find_all("tr"):
                cells = [c.get_text(strip=True) for c in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            if rows:
                rows_all.append(rows)
    except Exception:
        pass
    return rows_all


def _cells_rows_to_html(cells_rows: list[list[dict[str, Any]]], header_rows: int = 1, with_border: bool = True) -> str:
    """Build HTML table from cell dicts preserving colspan/rowspan and headers.

    Each cell dict: {"text": str, "colspan": int, "rowspan": int, "header": bool}
    """
    def tag_for(cell: dict[str, Any]) -> str:
        return "th" if (cell.get("header") or False) else "td"

    attrs_table = " border=\"1\"" if with_border else ""
    thead_html = ""
    tbody_html = ""
    rows_iter = enumerate(cells_rows)
    for idx, row in rows_iter:
        cells_html: list[str] = []
        for cell in row:
            name = tag_for(cell)
            colspan = int(cell.get("colspan", 1) or 1)
            rowspan = int(cell.get("rowspan", 1) or 1)
            parts = [name]
            if colspan > 1:
                parts.append(f"colspan=\"{colspan}\"")
            if rowspan > 1:
                parts.append(f"rowspan=\"{rowspan}\"")
            attrs = " " + " ".join(parts[1:]) if len(parts) > 1 else ""
            text = (cell.get("text") or "").strip()
            cells_html.append(f"<{name}{attrs}>{text}</{name}>")
        row_html = "<tr>" + "".join(cells_html) + "</tr>"
        if idx < max(0, header_rows):
            thead_html += row_html
        else:
            tbody_html += row_html
    if thead_html:
        return f"<table{attrs_table}><thead>{thead_html}</thead><tbody>{tbody_html}</tbody></table>"
    return f"<table{attrs_table}>{tbody_html}</table>"


def _cells_rows_to_plain_grid(cells_rows: list[list[dict[str, Any]]]) -> list[list[str]]:
    """Expand cells with rowspan/colspan into rectangular grid of strings.

    The text is placed in the top-left cell; spanned cells become empty strings.
    """
    grid: list[list[str]] = []
    carries: list[int] = []  # rows remaining for rowspan in each column

    def ensure_len(arr: list[int], n: int) -> None:
        if len(arr) < n:
            arr.extend([0] * (n - len(arr)))

    for row in cells_rows:
        row_vals: list[str] = []
        # Pre-fill carried columns with blanks as we advance through columns when placing cells
        # Place each new cell at next free column (carry == 0)
        for cell in row:
            # advance to next free column
            c = len(row_vals)
            while c < len(carries) and carries[c] > 0:
                row_vals.append("")
                c += 1
            # place this cell
            text = (cell.get("text") or "").strip()
            colspan = max(1, int(cell.get("colspan", 1) or 1))
            rowspan = max(1, int(cell.get("rowspan", 1) or 1))
            row_vals.append(text)
            for _ in range(colspan - 1):
                row_vals.append("")
            ensure_len(carries, len(row_vals))
            # mark carries for this span (downward rows)
            if rowspan > 1:
                for pos in range(c, min(c + colspan, len(carries))):
                    carries[pos] += (rowspan - 1)
        # After placing all cells, pad with blanks for any trailing carries
        # Consume remaining carries positions at the end of the row
        i = len(row_vals)
        while i < len(carries) and carries[i] > 0:
            row_vals.append("")
            i += 1
        # Decrement carries for next row
        carries = [max(0, v - 1) for v in carries]
        grid.append(row_vals)
    # Normalize width
    width = max((len(r) for r in grid), default=0)
    for r in grid:
        if len(r) < width:
            r.extend([""] * (width - len(r)))
    return grid


def _parse_tables_from_html(raw_html: str) -> list[dict[str, Any]]:
    """Parse HTML and return list of dicts: {html: str, rows_plain: list[list[str]]}.

    Preserves colspan/rowspan and detects header rows using <th> or first row.
    """
    out: list[dict[str, Any]] = []
    try:
        soup = BeautifulSoup(raw_html, "lxml")
        for tbl in soup.find_all("table"):
            cells_rows: list[list[dict[str, Any]]] = []
            header_rows = 0
            for tr in tbl.find_all("tr"):
                row_cells: list[dict[str, Any]] = []
                is_header_row = False
                for c in tr.find_all(["th", "td"]):
                    name = c.name.lower()
                    is_header = name == "th"
                    if is_header:
                        is_header_row = True
                    txt = c.get_text(strip=True)
                    colspan = int(c.get("colspan") or 1)
                    rowspan = int(c.get("rowspan") or 1)
                    row_cells.append({
                        "text": ftfy.fix_text(txt).strip(),
                        "colspan": colspan,
                        "rowspan": rowspan,
                        "header": is_header,
                    })
                if row_cells:
                    if is_header_row and header_rows == 0:
                        header_rows = 1
                    cells_rows.append(row_cells)
            if not cells_rows:
                continue
            html = _cells_rows_to_html(cells_rows, header_rows=header_rows or 1, with_border=True)
            grid = _cells_rows_to_plain_grid(cells_rows)
            out.append({
                "html": html,
                "rows_plain": grid,
            })
    except Exception:
        pass
    return out


def _extract_docx_tables_rows(path: Path) -> list[list[list[str]]]:
    out: list[list[list[str]]] = []
    try:
        import docx  # type: ignore
        document = docx.Document(str(path))
        for tbl in document.tables:
            rows: list[list[str]] = []
            for row in tbl.rows:
                cells = [ftfy.fix_text(cell.text or "").strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                out.append(rows)
    except Exception:
        pass
    return out


def _extract_odt_tables_cells_from_xml(xml: str) -> list[list[list[dict[str, Any]]]]:
    """Return per-table cells with spans preserved for ODT content.xml.

    table:table-cell may have table:number-columns-spanned / table:number-rows-spanned.
    """
    tables: list[list[list[dict[str, Any]]]] = []
    try:
        soup = BeautifulSoup(xml, "lxml")
        for tbl in soup.find_all(lambda tag: isinstance(tag.name, str) and tag.name.endswith("table")):
            table_rows: list[list[dict[str, Any]]] = []
            for tr in tbl.find_all(lambda tag: isinstance(tag.name, str) and tag.name.endswith("table-row")):
                row_cells: list[dict[str, Any]] = []
                cells = tr.find_all(lambda tag: isinstance(tag.name, str) and tag.name.endswith("table-cell"))
                for c in cells:
                    txt = c.get_text(strip=True)
                    colspan = int(c.get("table:number-columns-spanned") or 1)
                    rowspan = int(c.get("table:number-rows-spanned") or 1)
                    row_cells.append({
                        "text": ftfy.fix_text(txt).strip(),
                        "colspan": colspan,
                        "rowspan": rowspan,
                        "header": False,
                    })
                if row_cells:
                    table_rows.append(row_cells)
            if table_rows:
                tables.append(table_rows)
    except Exception:
        pass
    return tables


def _extract_delimited_table_rows_from_text(text: str) -> list[list[list[str]]]:
    rows_all: list[list[list[str]]] = []
    try:
        lines = [ln.strip() for ln in text.splitlines()]
        current: list[list[str]] = []
        for ln in lines:
            if (ln.count('|') >= 2) or (ln.count('\t') >= 1):
                parts = [p.strip() for p in re.split(r"\||\t", ln) if p.strip()]
                if len(parts) >= 2:
                    current.append(parts)
                    continue
            if current:
                rows_all.append(current)
                current = []
        if current:
            rows_all.append(current)
    except Exception:
        pass
    return rows_all


def _extract_whitespace_table_rows(text: str) -> list[list[list[str]]]:
    """Heuristic: detect fixed-width tables separated by 2+ spaces.

    Collect contiguous blocks with >=2 columns on consecutive lines.
    Skip border-only lines.
    """
    rows_all: list[list[list[str]]] = []
    try:
        lines = [ln.rstrip() for ln in text.splitlines()]
        current: list[list[str]] = []
        for ln in lines:
            # Skip border-like lines
            if re.fullmatch(r"[ \-\+=\|_\.]{5,}", ln.strip()):
                if current:
                    if len(current) >= 2:
                        rows_all.append(current)
                    current = []
                continue
            # Split by 2+ spaces
            parts = [p.strip() for p in re.split(r"\s{2,}", ln) if p.strip()]
            if len(parts) >= 2:
                current.append(parts)
            else:
                if current:
                    if len(current) >= 2:
                        rows_all.append(current)
                    current = []
        if current and len(current) >= 2:
            rows_all.append(current)
    except Exception:
        pass
    return rows_all


def _extract_key_fields_formal_doc(text: str) -> dict[str, str]:
    fields: dict[str, str] = {}
    try:
        def grab(pats: list[str]) -> str:
            for pat in pats:
                m = re.search(pat, text, re.IGNORECASE)
                if m:
                    return (m.group(1) or '').strip()
            return ''

        fields["work_permit_no"] = grab([r"Work\s*Permit\s*No\.?[:\-]?\s*([A-Za-z0-9\-\/]+)"])
        fields["visa_grant_number"] = grab([r"Visa\s+grant\s+number\s*[:\-]?\s*([A-Za-z0-9]+)"])
        fields["name"] = grab([r"Name\s*[:\-]?\s*([A-Z .'-]+)", r"Name\s+([A-Z .'-]+)"])
        fields["dob"] = grab([r"Date\s*of\s*Birth\s*[:\-]?\s*([0-9]{1,2}\s*[A-Za-z]{3,}\s*[0-9]{2,4}|[0-9]{2}\/[0-9]{2}\/[0-9]{2,4}|[0-9]{1,2}[A-Z]{3}[0-9]{2,4})"])
        fields["nationality"] = grab([r"Nationality\s*[:\-]?\s*([A-Za-z ]+)"])
        fields["passport_no"] = grab([r"Passport(?:\/|\s*or\s*Travel\s*Document)?\s*No\.?[:\-]?\s*([A-Za-z0-9]+)"])
        fields["employer"] = grab([r"Employer\s*[:\-]?\s*([A-Z0-9 ()&.,'-]+)", r"Name of the Employer\s*([A-Z0-9 ()&.,'-]+)"])
        fields["position"] = grab([r"(TECHNICAL\s+SUPERVISOR.*|Supervisor.*|Engineer.*|Manager.*)"])
        fields["date_of_issue"] = grab([r"Date\s*of\s*issue\s*[:\-]?\s*([0-9]{1,2}\s*[A-Za-z]{3,}\s*[0-9]{2,4})"])
        fields["date_of_expiry"] = grab([r"Date\s*of\s*Expiry\s*[:\-]?\s*([0-9]{1,2}\s*[A-Za-z]{3,}\s*[0-9]{2,4})"])
        # Drop empty
        fields = {k: v for k, v in fields.items() if v}
    except Exception:
        fields = {}
    return fields


def _extract_pdf_tables_html(path: Path) -> list[str]:
    html_tables: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for tbl in tables or []:
                    # Render minimal HTML table
                    rows = [
                        "<tr>" + "".join(f"<td>{(cell or '').strip()}</td>" for cell in row) + "</tr>"
                        for row in tbl
                    ]
                    html = "<table>" + "".join(rows) + "</table>"
                    html_tables.append(html)
    except Exception:
        pass
    return html_tables


def _read_pdf_text_plumber(path: Path) -> str:
    try:
        texts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                texts.append(t)
        return "\n".join(texts).strip()
    except Exception:
        return ""


def _read_pdf_text_hybrid(path: Path) -> str:
    """Extract text per-page; if a page has no text, OCR just that page.

    Respects NC_OCR_PDF_PAGE_LIMIT for OCR part.
    """
    settings = get_settings()
    ocr_limit = max(0, settings.ocr_pdf_page_limit or 0)
    texts: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            total_pages = len(pdf.pages)
            for idx, page in enumerate(pdf.pages, start=1):
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                if t.strip():
                    texts.append(t.strip())
                    continue
                # OCR this page if within limit
                if ocr_limit and idx > ocr_limit:
                    texts.append("")
                    continue
                try:
                    images = convert_from_path(
                        str(path), dpi=300, first_page=idx, last_page=idx
                    )
                    if images:
                        img = images[0]
                        # Reuse image OCR pipeline by passing through PIL path is not needed
                        # Apply same preprocessing and configs
                        # dump into file_id folder
                        try:
                            file_id_part = path.parent.name
                            prefix = f"{file_id_part}/{path.stem}_p{idx}"
                        except Exception:
                            prefix = f"{path.stem}_p{idx}"
                        text = _ocr_from_pil_image(img, dump_prefix=prefix)
                        texts.append(text)
                    else:
                        texts.append("")
                except Exception:
                    texts.append("")
    except Exception:
        return ""
    return "\n".join(texts).strip()


def _extract_pdf_images_ocr(path: Path, max_images: int = 10) -> list[str]:
    texts: list[str] = []
    try:
        reader = PdfReader(str(path))
        count = 0
        for page in reader.pages:
            try:
                xobjs = page.images  # type: ignore[attr-defined]
            except Exception:
                xobjs = []
            for img in xobjs:
                if count >= max_images:
                    break
                try:
                    from io import BytesIO

                    with Image.open(BytesIO(img.data)) as im:  # type: ignore[attr-defined]
                        # preprocess similar to image pipeline
                        try:
                            file_id_part = path.parent.name
                            prefix = f"{file_id_part}/{path.stem}_img{count}"
                        except Exception:
                            prefix = f"{path.stem}_img{count}"
                        t = _ocr_from_pil_image(im, dump_prefix=prefix)
                        if t:
                            texts.append(t)
                            count += 1
                except Exception:
                    continue
            if count >= max_images:
                break
    except Exception:
        return []
    return texts


def _cv2_preprocess_for_ocr(pil_img: Image.Image, dump_prefix: str | None = None) -> list[Image.Image]:
    settings = get_settings()
    variants: list[Image.Image] = []
    # Convert PIL -> OpenCV
    img = np.array(pil_img.convert("RGB"))
    img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
    # Deskew via moments angle
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.fastNlMeansDenoising(gray, h=10)
    # Try to estimate skew and rotate
    try:
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)
        lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=100, minLineLength=100, maxLineGap=10)
        angles: list[float] = []
        if lines is not None:
            for line in lines[:200]:
                x1, y1, x2, y2 = line[0]
                angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
                # Consider near-horizontal text lines
                if -45 < angle < 45:
                    angles.append(angle)
        if angles:
            median_angle = float(np.median(angles))
            if abs(median_angle) > 0.5:
                (h, w) = gray.shape[:2]
                M = cv2.getRotationMatrix2D((w // 2, h // 2), median_angle, 1.0)
                gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_LINEAR, borderMode=cv2.BORDER_REPLICATE)
    except Exception:
        pass
    # Thresholds
    th_otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
    th_mean = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 31, 10)
    th_gauss = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10)
    # Morphology to connect text
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
    morphed = cv2.morphologyEx(th_otsu, cv2.MORPH_OPEN, kernel)
    # Collect variants
    cv_variants = [gray, th_otsu, th_mean, th_gauss, morphed]
    for idx, arr in enumerate(cv_variants):
        try:
            im = Image.fromarray(arr)
            variants.append(im)
            if dump_prefix and (settings.ocr_debug_dump or str(settings.log_level).upper() == "DEBUG"):
                dump_dir = get_settings().data_dir / "artifacts"
                # Save via proper encoding
                out_path = dump_dir / f"{dump_prefix}_cv2_{idx}.png"
                out_path.parent.mkdir(parents=True, exist_ok=True)
                cv2.imwrite(str(out_path), arr)
                try:
                    logger.debug("ocr_dump_saved", path=str(out_path))
                except Exception:
                    pass
        except Exception:
            continue
    return variants


def _save_pil_variant(img: Image.Image, prefix: str, idx: int) -> None:
    try:
        settings = get_settings()
        if not (settings.ocr_debug_dump or str(settings.log_level).upper() == "DEBUG"):
            return
        dump_dir = settings.data_dir / "artifacts"
        dump_dir.mkdir(parents=True, exist_ok=True)
        out_path = dump_dir / f"{prefix}_pil_{idx}.png"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            logger.debug("ocr_dump_try", path=str(out_path))
        except Exception:
            pass
        try:
            img.save(out_path)
            try:
                logger.debug("ocr_dump_saved", path=str(out_path))
            except Exception:
                pass
        except Exception as e:
            try:
                logger.warning("ocr_dump_error", path=str(out_path), error=str(e))
            except Exception:
                pass
    except Exception as e:
        try:
            logger.warning("ocr_dump_error_setup", error=str(e))
        except Exception:
            pass
        return


def _ocr_from_pil_image(pil_img: Image.Image, dump_prefix: str | None = None) -> str:
    settings = get_settings()
    variants: list[Image.Image] = []
    try:
        base = pil_img.convert("RGB")
        # Force dump original
        try:
            if dump_prefix and (settings.ocr_debug_dump or str(settings.log_level).upper() == "DEBUG"):
                dump_dir = settings.data_dir / "artifacts"
                out_path = dump_dir / f"{dump_prefix}_orig.png"
                out_path.parent.mkdir(parents=True, exist_ok=True)
                try:
                    logger.debug("ocr_dump_try", path=str(out_path))
                except Exception:
                    pass
                try:
                    base.save(out_path)
                    logger.debug("ocr_dump_saved", path=str(out_path))
                except Exception as e:
                    logger.warning("ocr_dump_error", path=str(out_path), error=str(e))
        except Exception as e:
            try:
                logger.warning("ocr_dump_error_setup", error=str(e))
            except Exception:
                pass
        mx = max(base.size)
        if mx < 800:
            base = base.resize((base.width * 3, base.height * 3), Image.LANCZOS)
        elif mx < 1200:
            base = base.resize((base.width * 2, base.height * 2), Image.LANCZOS)
        g = base.convert("L")
        pil_variants = [
            g,
            ImageOps.autocontrast(g),
            ImageEnhance.Contrast(g).enhance(1.8),
            ImageOps.invert(g),
            g.filter(ImageFilter.UnsharpMask(radius=2, percent=150)),
        ]
        for i, v in enumerate(pil_variants):
            variants.append(v)
            if dump_prefix:
                _save_pil_variant(v, dump_prefix, i)
        # OpenCV variants
        variants.extend(_cv2_preprocess_for_ocr(base, dump_prefix=dump_prefix))
    except Exception:
        variants = [pil_img]
    configs = [
        f"--oem 1 --psm {settings.ocr_tesseract_psm}",
        "--oem 3 --psm 6",
        "--oem 3 --psm 4",
        "--oem 3 --psm 7",
        "--oem 1 --psm 11",
        "--oem 1 --psm 12",
        "--oem 1 --psm 13",
        "--oem 1 --psm 6 -c tessedit_char_blacklist=~`^*_{}[]|\\",
        "--oem 1 --psm 6 -c preserve_interword_spaces=1",
    ]
    try:
        logger.info(
            "ocr_attempt_start",
            variants=len(variants),
            dump_enabled=settings.ocr_debug_dump,
            prefix=dump_prefix,
            ocr_langs=settings.ocr_langs,
        )
    except Exception:
        pass
    for im in variants:
        for cfg in configs:
            try:
                text = ocr_image_to_string(im, lang=settings.ocr_langs, config=cfg).strip()
                if text:
                    try:
                        logger.info("ocr_attempt_ok", length=len(text), config=cfg)
                    except Exception:
                        pass
                    return text
            except Exception:
                continue
    try:
        logger.info("ocr_attempt_empty")
    except Exception:
        pass
    return ""


def _read_image_text(path: Path) -> str:
    settings = get_settings()
    with Image.open(path) as img:
        # Put dumps under file_id folder if possible
        try:
            file_id_part = path.parent.name
            dump_prefix = f"{file_id_part}/{path.stem}"
        except Exception:
            dump_prefix = path.stem
        return _ocr_from_pil_image(img, dump_prefix=dump_prefix)


def _pdf_has_text_layer(path: Path) -> bool:
    try:
        txt = pdf_extract_text(str(path))
        return bool(txt and txt.strip())
    except Exception:
        return False


def _ocr_pdf_pages_to_text(path: Path, dpi: int = 300) -> str:
    settings = get_settings()
    limit = settings.ocr_pdf_page_limit
    # Use first_page/last_page to avoid rendering all pages
    images = convert_from_path(str(path), dpi=dpi, first_page=1, last_page=limit if limit else None)
    texts: list[str] = []
    for img in images:
        texts.append(ocr_image_to_string(img, lang=settings.ocr_langs, config=f"--psm {settings.ocr_tesseract_psm}"))
    return "\n".join(texts)


def _extract_pdf_tables_rows(path: Path) -> list[list[list[str]]]:
    tables_rows: list[list[list[str]]] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for tbl in tables or []:
                    rows: list[list[str]] = []
                    for row in tbl:
                        cells = [(cell or "").strip() for cell in row]
                        rows.append(cells)
                    if rows:
                        tables_rows.append(rows)
    except Exception:
        pass
    return tables_rows


def _render_html_table(rows: list[list[str]]) -> str:
    tr_list = [
        "<tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>"
        for row in rows
    ]
    return "<table>" + "".join(tr_list) + "</table>"


def _render_plain_table(rows: list[list[str]]) -> str:
    return "\n".join(" | ".join(row) for row in rows)


def _read_docx_text(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception:
        return ""
    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_docx_images_ocr(path: Path) -> list[str]:
    try:
        import docx  # type: ignore
    except Exception:
        return []
    texts: list[str] = []
    document = docx.Document(str(path))
    # Iterate related parts to find images
    try:
        rel_parts = document.part.related_parts
        for rel_id, part in rel_parts.items():  # type: ignore[assignment]
            try:
                if getattr(part, "content_type", "").startswith("image/"):
                    from io import BytesIO

                    with Image.open(BytesIO(part.blob)) as img:  # type: ignore[attr-defined]
                        # Reuse image OCR pipeline
                        tmp_path = path  # not used; call OCR directly on PIL image
                        # Preprocess similar to _read_image_text
                        im = img
                        try:
                            if max(im.size) < 1200:
                                im = im.resize((im.width * 2, im.height * 2), Image.LANCZOS)
                            im = im.convert("L")
                            im = im.point(lambda x: 0 if x < 140 else 255, "1")
                        except Exception:
                            pass
                        configs = ["--oem 1 --psm 6", "--oem 1 --psm 3", "--oem 1 --psm 11"]
                        for cfg in configs:
                            t = ocr_image_to_string(im, lang=get_settings().ocr_langs, config=cfg).strip()
                            if t:
                                texts.append(t)
                                break
            except Exception:
                continue
    except Exception:
        pass
    return texts


def parse_document_to_text(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()

    def _detect_type(p: Path) -> str:
        try:
            with p.open("rb") as f:
                head = f.read(4096)
            h = head[:16]
            if head.startswith(b"%PDF"):
                return "pdf"
            if head.startswith(b"\x89PNG\r\n\x1a\n"):
                return "png"
            if head.startswith(b"\xFF\xD8\xFF"):
                return "jpg"
            if head[:2] == b"PK":
                # Could be DOCX/ZIP/ODT; cheap check — rely on suffix next
                return "zip"
            if head[:8] == b"{\\rtf1":
                return "rtf"
            txt_sample = head.decode("utf-8", errors="ignore").lower()
            if "<html" in txt_sample or "<!doctype html" in txt_sample:
                return "html"
            if "," in txt_sample and "\n" in txt_sample:
                return "csv"
            return "txt"
        except Exception:
            return "txt"

    ftype = _detect_type(path)
    try:
        logger.info("detect_file_type", suffix=suffix, ftype=ftype, path=str(path))
    except Exception:
        pass

    timings: dict[str, float] = {}
    metrics: dict[str, Any] = {}
    t_step = time.perf_counter()
    if ftype == "txt" or (suffix in {".txt", ""} and ftype is None):
        text = _read_text_file(path)
        timings["txt_read_ms"] = (time.perf_counter() - t_step) * 1000
    elif ftype == "pdf" or suffix == ".pdf":
        t_pdf = time.perf_counter()
        if not _pdf_quick_sanity(path):
            timings["pdf_sanity_ms"] = (time.perf_counter() - t_pdf) * 1000
            return ParsedDocument(full_text="", pages=[])
        timings["pdf_sanity_ms"] = (time.perf_counter() - t_pdf) * 1000
        t_layer = time.perf_counter()
        has_text_layer = _pdf_has_text_layer(path)
        timings["pdf_text_layer_check_ms"] = (time.perf_counter() - t_layer) * 1000
        if has_text_layer:
            # Hybrid: try per-page; OCR only empty pages
            t_txt = time.perf_counter()
            text = _read_pdf_text_hybrid(path) or _read_pdf_text_plumber(path) or _read_pdf_text(path)
            timings["pdf_text_extract_ms"] = (time.perf_counter() - t_txt) * 1000
            if not text:
                # Fallback to OCR for tricky text-layer PDFs
                t_ocr = time.perf_counter()
                text = _ocr_pdf_pages_to_text(path)
                timings["pdf_ocr_pages_ms"] = (time.perf_counter() - t_ocr) * 1000
        else:
            # Guard rails for OCR on big docs
            settings = get_settings()
            size_mb = path.stat().st_size / (1024 * 1024)
            try:
                with pdfplumber.open(str(path)) as pdf:
                    num_pages = len(pdf.pages)
            except Exception:
                num_pages = 0
            if size_mb > settings.ocr_pdf_max_mb or (settings.ocr_pdf_max_pages and num_pages > settings.ocr_pdf_max_pages):
                text = ""  # skip OCR
            else:
                t_ocr = time.perf_counter()
                text = _ocr_pdf_pages_to_text(path)
                timings["pdf_ocr_pages_ms"] = (time.perf_counter() - t_ocr) * 1000
        t_tbl = time.perf_counter()
        tables = _extract_pdf_tables_rows(path)
        timings["pdf_tables_ms"] = (time.perf_counter() - t_tbl) * 1000
        pages: list[dict[str, Any]] = [{"index": 0, "text": text}]
        # OCR embedded images if text is still weak
        t_img = time.perf_counter()
        image_texts = _extract_pdf_images_ocr(path)
        timings["pdf_image_ocr_ms"] = (time.perf_counter() - t_img) * 1000
        if image_texts:
            pages.append({
                "index": len(pages),
                "text": "\n\n".join(image_texts),
                "elements": [{"type": "image_ocr", "description": t} for t in image_texts],
            })
            text = (text + "\n\n" + "\n\n".join(image_texts)).strip()
        # Optional captioning for embedded images (gated by flag) — batch with caching and heuristics
        try:
            if get_settings().captioning_enabled:
                settings = get_settings()
                t_cap = time.perf_counter()
                images_for_caption: list[Image.Image] = []
                try:
                    reader = PdfReader(str(path))
                    for page in reader.pages:
                        imgs = getattr(page, "images", []) or []  # type: ignore[attr-defined]
                        for img in imgs:
                            from io import BytesIO
                            try:
                                bio = BytesIO(img.data)  # type: ignore[attr-defined]
                                im = Image.open(bio)
                                im.load()
                            except Exception:
                                continue
                            if max(im.size) < max(1, settings.caption_min_image_px):
                                continue
                            # Heuristics: skip extreme aspect ratios and very low-entropy images
                            try:
                                w, h = im.size
                                aspect = (w / max(1, h)) if h > 0 else 999.0
                                aspect = max(aspect, 1.0 / max(1e-6, aspect))  # unify ratio > 1
                                if aspect > max(1.0, settings.caption_max_aspect_ratio):
                                    continue
                                # entropy proxy: histogram dispersion
                                hist = im.convert("L").histogram()
                                total = float(sum(hist)) or 1.0
                                import math
                                probs = [v / total for v in hist if v > 0]
                                ent = -sum(p * math.log(p + 1e-12) for p in probs)
                                if ent < max(0.0, settings.caption_min_entropy):
                                    continue
                            except Exception:
                                pass
                            images_for_caption.append(im)
                            if len(images_for_caption) >= max(1, settings.caption_max_images_per_doc):
                                break
                        if len(images_for_caption) >= max(1, settings.caption_max_images_per_doc):
                            break
                except Exception:
                    images_for_caption = []
                cap_texts: list[str] = []
                if images_for_caption:
                    caps, cap_metrics = caption_images_with_cache(images_for_caption)
                    cap_texts = [c.text for c in caps if c.text]
                    try:
                        metrics["caption"] = {"count": len(caps), **cap_metrics}
                    except Exception:
                        pass
                timings["pdf_caption_ms"] = (time.perf_counter() - t_cap) * 1000
                if cap_texts:
                    pages.append({
                        "index": len(pages),
                        "text": "\n\n".join(cap_texts),
                        "elements": [{"type": "image_caption", "description": c.text, "model": c.model} for c in caps if c.text],
                    })
        except Exception:
            pass
        if tables:
            tables_html = [_render_html_table(rows) for rows in tables]
            tables_plain = [_render_plain_table(rows) for rows in tables]
            # Add tables page with HTML elements and plain text content
            elements = [{"type": "table_html", "description": html} for html in tables_html]
            pages.append({"index": 1, "text": "\n\n".join(tables_plain), "elements": elements})
            # Concatenate plain tables to full text for searchability
            text = (text + "\n\n" + "\n\n".join(tables_plain)).strip()
        # Extract key fields
        t_fields = time.perf_counter()
        fields = _extract_key_fields_formal_doc(text)
        if fields:
            pages.append({"index": len(pages), "text": "", "elements": [{"type": "fields", "description": json.dumps(fields, ensure_ascii=False)}]})
        timings["fields_extract_ms"] = timings.get("fields_extract_ms", 0.0) + (time.perf_counter() - t_fields) * 1000
        t_norm = time.perf_counter()
        text = _normalize_output_text(text)
        for p in pages:
            p["text"] = _normalize_output_text(p.get("text", ""))
        timings["normalize_ms"] = (time.perf_counter() - t_norm) * 1000
        return ParsedDocument(full_text=text, pages=pages, timings_ms=timings, metrics=(metrics or None))
    elif ftype in {"png", "jpg"} or suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}:
        t_img = time.perf_counter()
        text = _normalize_output_text(_read_image_text(path))
        timings["image_ocr_ms"] = (time.perf_counter() - t_img) * 1000
    elif ftype == "docx" or suffix in {".docx"}:
        t_docx = time.perf_counter()
        text = _read_docx_text(path)
        timings["docx_text_ms"] = (time.perf_counter() - t_docx) * 1000
        # OCR for embedded images
        t_img = time.perf_counter()
        image_texts = _extract_docx_images_ocr(path)
        timings["docx_images_ocr_ms"] = (time.perf_counter() - t_img) * 1000
        # Extract tables
        t_tbl = time.perf_counter()
        docx_tables = _extract_docx_tables_rows(path)
        timings["docx_tables_ms"] = (time.perf_counter() - t_tbl) * 1000
        tables_html = [_render_html_table(rows) for rows in docx_tables]
        tables_plain = [_render_plain_table(rows) for rows in docx_tables]
        pages: list[dict[str, Any]] = []
        if text:
            pages.append({"index": 0, "text": text})
        if image_texts:
            pages.append({
                "index": len(pages),
                "text": "\n\n".join(image_texts),
                "elements": [{"type": "image_ocr", "description": t} for t in image_texts],
            })
            text = (text + "\n\n" + "\n\n".join(image_texts)).strip()
        # Optional captioning for embedded images in DOCX — batch with caching and heuristics
        try:
            if get_settings().captioning_enabled:
                settings = get_settings()
                t_cap = time.perf_counter()
                images_for_caption: list[Image.Image] = []
                try:
                    import docx  # type: ignore
                    document = docx.Document(str(path))
                    rel_parts = document.part.related_parts
                    for rel_id, part in rel_parts.items():  # type: ignore[assignment]
                        if getattr(part, "content_type", "").startswith("image/"):
                            from io import BytesIO
                            try:
                                bio = BytesIO(part.blob)  # type: ignore[attr-defined]
                                img = Image.open(bio)
                                img.load()
                            except Exception:
                                continue
                            if max(img.size) < max(1, settings.caption_min_image_px):
                                continue
                            try:
                                w, h = img.size
                                aspect = (w / max(1, h)) if h > 0 else 999.0
                                aspect = max(aspect, 1.0 / max(1e-6, aspect))
                                if aspect > max(1.0, settings.caption_max_aspect_ratio):
                                    continue
                                hist = img.convert("L").histogram()
                                total = float(sum(hist)) or 1.0
                                import math
                                probs = [v / total for v in hist if v > 0]
                                ent = -sum(p * math.log(p + 1e-12) for p in probs)
                                if ent < max(0.0, settings.caption_min_entropy):
                                    continue
                            except Exception:
                                pass
                            images_for_caption.append(img)
                            if len(images_for_caption) >= max(1, settings.caption_max_images_per_doc):
                                break
                except Exception:
                    images_for_caption = []
                cap_texts: list[str] = []
                if images_for_caption:
                    caps, cap_metrics = caption_images_with_cache(images_for_caption)
                    cap_texts = [c.text for c in caps if c.text]
                    try:
                        metrics["caption"] = {"count": len(caps), **cap_metrics}
                    except Exception:
                        pass
                timings["docx_caption_ms"] = (time.perf_counter() - t_cap) * 1000
                if cap_texts:
                    pages.append({
                        "index": len(pages),
                        "text": "\n\n".join(cap_texts),
                        "elements": [{"type": "image_caption", "description": c.text, "model": c.model} for c in caps if c.text],
                    })
        except Exception:
            pass
        if tables_plain:
            pages.append({
                "index": len(pages),
                "text": "\n\n".join(tables_plain),
                "elements": [{"type": "table_html", "description": html} for html in tables_html],
            })
            text = (text + "\n\n" + "\n\n".join(tables_plain)).strip()
        # Extract key fields
        t_fields = time.perf_counter()
        fields = _extract_key_fields_formal_doc(text)
        if fields:
            pages.append({"index": len(pages), "text": "", "elements": [{"type": "fields", "description": json.dumps(fields, ensure_ascii=False)}]})
        timings["fields_extract_ms"] = timings.get("fields_extract_ms", 0.0) + (time.perf_counter() - t_fields) * 1000
        t_norm = time.perf_counter()
        text = _normalize_output_text(text)
        for p in pages:
            p["text"] = _normalize_output_text(p.get("text", ""))
        timings["normalize_ms"] = (time.perf_counter() - t_norm) * 1000
        return ParsedDocument(full_text=text, pages=pages if pages else [], timings_ms=timings, metrics=(metrics or None))
    elif suffix == ".doc":
        t_doc = time.perf_counter()
        text = _normalize_output_text(_read_doc_binary_text(path))
        timings["doc_text_ms"] = (time.perf_counter() - t_doc) * 1000
        # Try to extract simple delimited tables from text
        t_tbl = time.perf_counter()
        tables_rows = _extract_delimited_table_rows_from_text(text)
        # Also try whitespace-separated columns
        tables_rows += _extract_whitespace_table_rows(text)
        timings["doc_tables_ms"] = (time.perf_counter() - t_tbl) * 1000
        tables_html = [_render_html_table(rows) for rows in tables_rows]
        tables_plain = [_render_plain_table(rows) for rows in tables_rows]
        pages = [{"index": 0, "text": text}]
        if tables_plain:
            pages.append({
                "index": 1,
                "text": "\n\n".join(tables_plain),
                "elements": [{"type": "table_html", "description": html} for html in tables_html],
            })
            text = (text + "\n\n" + "\n\n".join(tables_plain)).strip()
        # Extract known fields
        t_fields = time.perf_counter()
        fields = _extract_key_fields_formal_doc(text)
        if fields:
            pages.append({"index": len(pages), "text": "", "elements": [{"type": "fields", "description": json.dumps(fields, ensure_ascii=False)}]})
        timings["fields_extract_ms"] = timings.get("fields_extract_ms", 0.0) + (time.perf_counter() - t_fields) * 1000
        t_norm = time.perf_counter()
        text = _normalize_output_text(text)
        for p in pages:
            p["text"] = _normalize_output_text(p.get("text", ""))
        timings["normalize_ms"] = (time.perf_counter() - t_norm) * 1000
        return ParsedDocument(full_text=text, pages=pages, timings_ms=timings, metrics=(metrics or None))
    elif ftype == "rtf" or suffix == ".rtf":
        t_rtf = time.perf_counter()
        text = _normalize_output_text(_read_rtf_text(path))
        timings["rtf_text_ms"] = (time.perf_counter() - t_rtf) * 1000
        # Prefer HTML-rendered tables from unrtf to preserve spans/headers
        t_tbl = time.perf_counter()
        rtf_html = _rtf_to_html(path)
        tables_html_blocks: list[str] = []
        tables_plain_rows: list[list[list[str]]] = []
        if rtf_html:
            parsed = _parse_tables_from_html(rtf_html)
            for item in parsed:
                tables_html_blocks.append(item["html"])  # type: ignore[index]
                tables_plain_rows.append(item["rows_plain"])  # type: ignore[index]
        # Fallback heuristics if no HTML tables detected
        if not tables_plain_rows:
            heur_rows = _extract_delimited_table_rows_from_text(text)
            heur_rows += _extract_whitespace_table_rows(text)
            tables_plain_rows = heur_rows
            tables_html_blocks = [_render_html_table(rows) for rows in heur_rows]
        timings["rtf_tables_ms"] = (time.perf_counter() - t_tbl) * 1000
        tables_html = tables_html_blocks
        tables_plain = [_render_plain_table(rows) for rows in tables_plain_rows]
        pages = [{"index": 0, "text": text}]
        if tables_plain:
            pages.append({
                "index": 1,
                "text": "\n\n".join(tables_plain),
                "elements": [{"type": "table_html", "description": html} for html in tables_html],
            })
            text = (text + "\n\n" + "\n\n".join(tables_plain)).strip()
        t_fields = time.perf_counter()
        fields = _extract_key_fields_formal_doc(text)
        if fields:
            pages.append({"index": len(pages), "text": "", "elements": [{"type": "fields", "description": json.dumps(fields, ensure_ascii=False)}]})
        timings["fields_extract_ms"] = timings.get("fields_extract_ms", 0.0) + (time.perf_counter() - t_fields) * 1000
        t_norm = time.perf_counter()
        text = _normalize_output_text(text)
        for p in pages:
            p["text"] = _normalize_output_text(p.get("text", ""))
        timings["normalize_ms"] = (time.perf_counter() - t_norm) * 1000
        return ParsedDocument(full_text=text, pages=pages, timings_ms=timings, metrics=(metrics or None))
    elif suffix == ".odt":
        # Extract text and tables from content.xml
        t_odt = time.perf_counter()
        try:
            with zipfile.ZipFile(str(path), "r") as zf:
                xml = zf.open("content.xml").read().decode("utf-8", errors="ignore")
        except Exception:
            xml = ""
        text = _normalize_output_text(BeautifulSoup(xml, "lxml").get_text("\n") if xml else _read_odt_text(path))
        timings["odt_text_ms"] = (time.perf_counter() - t_odt) * 1000
        t_tbl = time.perf_counter()
        cells_tables = _extract_odt_tables_cells_from_xml(xml) if xml else []
        tables_html: list[str] = []
        tables_plain: list[str] = []
        if cells_tables:
            for cells_rows in cells_tables:
                html = _cells_rows_to_html(cells_rows, header_rows=1, with_border=True)
                grid = _cells_rows_to_plain_grid(cells_rows)
                tables_html.append(html)
                tables_plain.append(_render_plain_table(grid))
        if not tables_plain:
            heur_rows = _extract_whitespace_table_rows(text)
            tables_html = [_render_html_table(rows) for rows in heur_rows]
            tables_plain = [_render_plain_table(rows) for rows in heur_rows]
        timings["odt_tables_ms"] = (time.perf_counter() - t_tbl) * 1000
        pages = [{"index": 0, "text": text}]
        if tables_plain:
            pages.append({
                "index": 1,
                "text": "\n\n".join(tables_plain),
                "elements": [{"type": "table_html", "description": html} for html in tables_html],
            })
            text = (text + "\n\n" + "\n\n".join(tables_plain)).strip()
        t_fields = time.perf_counter()
        fields = _extract_key_fields_formal_doc(text)
        if fields:
            pages.append({"index": len(pages), "text": "", "elements": [{"type": "fields", "description": json.dumps(fields, ensure_ascii=False)}]})
        timings["fields_extract_ms"] = timings.get("fields_extract_ms", 0.0) + (time.perf_counter() - t_fields) * 1000
        t_norm = time.perf_counter()
        text = _normalize_output_text(text)
        for p in pages:
            p["text"] = _normalize_output_text(p.get("text", ""))
        timings["normalize_ms"] = (time.perf_counter() - t_norm) * 1000
        return ParsedDocument(full_text=text, pages=pages, timings_ms=timings, metrics=(metrics or None))
    elif ftype == "csv" or suffix in {".csv"}:
        # Parse CSV to HTML table and plain text
        try:
            # Detect encoding
            enc_text = None
            try:
                best = detect_encoding_from_path(str(path)).best()
                if best is not None:
                    enc_text = best.output_text(stripped=False)
            except Exception:
                enc_text = None
            rows: list[list[str]] = []
            if enc_text is not None:
                # Read from normalized text
                for line in ftfy.fix_text(enc_text).splitlines():
                    # Use csv reader on a list with single line to preserve parsing rules
                    for row in csv.reader([line]):
                        rows.append([col.strip(" \t\ufeff") for col in row])
            else:
                with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
                    reader = csv.reader(f)
                    for row in reader:
                        rows.append([ftfy.fix_text(col).strip(" \t\ufeff") for col in row])
            html = _render_html_table(rows)
            plain = _normalize_output_text(_render_plain_table(rows))
            pages = [
                {"index": 0, "text": plain, "elements": [{"type": "table_html", "description": html}]}
            ]
            timings["csv_parse_ms"] = (time.perf_counter() - t_step) * 1000
            return ParsedDocument(full_text=plain, pages=pages, timings_ms=timings, metrics=(metrics or None))
        except Exception:
            text = ""
    elif suffix in {".md", ".markdown"}:
        # Simple markdown strip: remove fenced code markers and headers
        t_md = time.perf_counter()
        raw = _read_text_file(path)
        text = raw.replace("```", "\n").replace("#", "").strip()
        timings["md_text_ms"] = (time.perf_counter() - t_md) * 1000
    elif ftype == "html" or suffix in {".html", ".htm"}:
        t_html = time.perf_counter()
        raw = _read_text_file(path)
        text = _extract_text_from_html(raw)
        timings["html_text_ms"] = (time.perf_counter() - t_html) * 1000
        t_tbl = time.perf_counter()
        html_tables_rows = _extract_html_tables_rows_from_html(raw)
        timings["html_tables_ms"] = (time.perf_counter() - t_tbl) * 1000
        tables_html = [_render_html_table(rows) for rows in html_tables_rows]
        tables_plain = [_render_plain_table(rows) for rows in html_tables_rows]
        pages = [{"index": 0, "text": text}]
        if tables_plain:
            pages.append({
                "index": 1,
                "text": "\n\n".join(tables_plain),
                "elements": [{"type": "table_html", "description": html} for html in tables_html],
            })
            text = (text + "\n\n" + "\n\n".join(tables_plain)).strip()
        t_fields = time.perf_counter()
        fields = _extract_key_fields_formal_doc(text)
        if fields:
            pages.append({"index": len(pages), "text": "", "elements": [{"type": "fields", "description": json.dumps(fields, ensure_ascii=False)}]})
        timings["fields_extract_ms"] = timings.get("fields_extract_ms", 0.0) + (time.perf_counter() - t_fields) * 1000
        return ParsedDocument(full_text=text, pages=pages, timings_ms=timings, metrics=(metrics or None))
    else:
        text = ""
    t_norm_end = time.perf_counter()
    text = _normalize_output_text(text)
    timings["normalize_ms"] = timings.get("normalize_ms", 0.0) + (time.perf_counter() - t_norm_end) * 1000
    pages = [{"index": 0, "text": text}] if text else []
    return ParsedDocument(full_text=text, pages=pages, timings_ms=timings, metrics=(metrics or None))


